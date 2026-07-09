"""
Feature 1 & 2 - extracts raw text from an uploaded plan file, regardless of
format, so the workout/nutrition parsers always operate on plain text. Every
extractor returns (text, warning|None) rather than raising for "no text
found", so the caller can surface a useful message instead of a 500.
"""
import io
import logging

from fastapi import UploadFile, HTTPException

logger = logging.getLogger("evofit.plan_import")

SUPPORTED_EXTENSIONS = {"txt", "pdf", "docx", "png", "jpg", "jpeg"}
MAX_FILE_SIZE_BYTES = 15 * 1024 * 1024  # 15 MB


def _extension(filename: str) -> str:
    if not filename or "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)

    text = "\n".join(text_parts).strip()
    if text:
        return text

    # Scanned/image-only PDF: fall back to OCR, page by page.
    return _extract_pdf_via_ocr(data)


def _extract_pdf_via_ocr(data: bytes) -> str:
    try:
        from pdf2image import convert_from_bytes
    except ImportError:
        logger.warning("pdf2image not available; cannot OCR a scanned PDF")
        return ""

    import pytesseract

    try:
        images = convert_from_bytes(data, dpi=200)
    except Exception:
        logger.exception("Failed to rasterize PDF for OCR fallback")
        return ""

    text_parts = [pytesseract.image_to_string(img) for img in images]
    return "\n".join(text_parts).strip()


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs if p.text.strip()]

    # Plans are frequently laid out in tables (day | exercise | sets | reps),
    # so table cells are pulled in too, one row per line.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def _extract_image(data: bytes) -> str:
    import pytesseract
    from PIL import Image, ImageOps

    image = Image.open(io.BytesIO(data))
    image = ImageOps.exif_transpose(image)
    if image.mode != "L":
        image = image.convert("L")
    return pytesseract.image_to_string(image)


def extract_text(file: UploadFile, data: bytes) -> tuple[str, str]:
    """Returns (raw_text, source_type). Raises HTTPException on bad input."""
    if len(data) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=400, detail="File is too large (max 15MB).")

    ext = _extension(file.filename or "")
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '.{ext}'. Supported: txt, pdf, docx, png, jpg, jpeg.",
        )

    try:
        if ext == "txt":
            text = _extract_txt(data)
        elif ext == "pdf":
            text = _extract_pdf(data)
            ext = "pdf"
        elif ext == "docx":
            text = _extract_docx(data)
        else:  # png / jpg / jpeg
            text = _extract_image(data)
            ext = "image"
    except HTTPException:
        raise
    except Exception:
        logger.exception("Failed to extract text from uploaded file %s", file.filename)
        raise HTTPException(status_code=422, detail="Could not read that file. It may be corrupted or unsupported.")

    text = (text or "").strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="No readable text was found in that file. Try a clearer scan/photo or paste the plan as text.",
        )

    source_type = "image" if ext in ("png", "jpg", "jpeg") else ext
    return text, source_type
