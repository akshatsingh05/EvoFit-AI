"""
Feature 8 - exports a comparison + AI analysis report as PDF or DOCX.
Scope is 'workout', 'nutrition', or 'combined' (both, requires both analyses
be passed in).
"""
import io
from datetime import datetime, timezone


def _title_for(scope: str) -> str:
    return {
        "workout": "Workout Plan Comparison Report",
        "nutrition": "Nutrition Plan Comparison Report",
        "combined": "Combined Plan Comparison Report",
    }.get(scope, "Plan Comparison Report")


def _metric_rows(comparison: dict) -> list[tuple[str, str, str]]:
    mine, evofit = comparison["mine"], comparison.get("evofit")
    rows = []
    if comparison["type"] == "workout":
        rows.append(("Training days/week", str(mine["workout_days_count"]),
                      str(evofit["workout_days_count"]) if evofit else "-"))
        rows.append(("Total exercises", str(mine["total_exercises"]), str(evofit["total_exercises"]) if evofit else "-"))
        rows.append(("Detected split", mine["detected_split"], "-"))
        rows.append(("Muscle balance score", f"{mine['muscle_balance_score']}/100",
                      f"{evofit['muscle_balance_score']}/100" if evofit else "-"))
        rows.append(("Effectiveness score", f"{mine['effectiveness_score']}/100",
                      f"{evofit['effectiveness_score']}/100" if evofit else "-"))
        for group, sets in mine["muscle_volume"].items():
            evo_sets = evofit["muscle_volume"].get(group) if evofit else None
            rows.append((f"{group.title()} volume (sets/wk)", str(sets), str(evo_sets) if evo_sets is not None else "-"))
    else:
        rows.append(("Avg daily calories", str(mine["avg_daily_calories"]),
                      str(evofit["avg_daily_calories"]) if evofit else "-"))
        rows.append(("Avg daily protein (g)", str(mine["avg_daily_protein_g"]),
                      str(evofit["avg_daily_protein_g"]) if evofit else "-"))
        rows.append(("Avg daily carbs (g)", str(mine["avg_daily_carbs_g"]),
                      str(evofit["avg_daily_carbs_g"]) if evofit else "-"))
        rows.append(("Avg daily fat (g)", str(mine["avg_daily_fat_g"]),
                      str(evofit["avg_daily_fat_g"]) if evofit else "-"))
        rows.append(("Avg water (ml)", str(mine["avg_water_ml"] or "-"),
                      str(evofit["avg_water_ml"]) if evofit else "-"))
        rows.append(("Unique foods", str(mine["unique_foods"]), str(evofit["unique_foods"]) if evofit else "-"))
        rows.append(("Effectiveness score", f"{mine['effectiveness_score']}/100",
                      f"{evofit['effectiveness_score']}/100" if evofit else "-"))
    return rows


def build_pdf_report(sections: list[dict]) -> bytes:
    """sections: list of {scope, plan_name, comparison, observations, suggestions}"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=colors.HexColor("#1B5E20"))
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=colors.HexColor("#2E7D32"))
    body = styles["BodyText"]

    story = [Paragraph("EvoFit AI - Plan Comparison Report", h1),
             Paragraph(datetime.now(timezone.utc).strftime("Generated %B %d, %Y"), body),
             Spacer(1, 0.2 * inch)]

    for section in sections:
        story.append(Paragraph(f"{section['plan_name']} ({section['scope'].title()})", h2))
        rows = [["Metric", "My Plan", "EvoFit AI Plan"]] + _metric_rows(section["comparison"])
        table = Table(rows, colWidths=[2.6 * inch, 1.8 * inch, 1.8 * inch])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E7D32")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F5F5")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.15 * inch))

        story.append(Paragraph("AI Analysis", styles["Heading3"]))
        for obs in section["observations"]:
            story.append(Paragraph(f"&bull; {obs}", body))
        story.append(Spacer(1, 0.1 * inch))

        if section["suggestions"]:
            story.append(Paragraph("Suggestions", styles["Heading3"]))
            for s in section["suggestions"]:
                story.append(Paragraph(f"<b>{s['title']}</b>: {s['detail']} <i>({s['reason']})</i>", body))
        story.append(Spacer(1, 0.25 * inch))

    doc.build(story)
    return buffer.getvalue()


def build_docx_report(sections: list[dict]) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor

    document = docx.Document()
    title = document.add_heading("EvoFit AI - Plan Comparison Report", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    document.add_paragraph(datetime.now(timezone.utc).strftime("Generated %B %d, %Y"))

    for section in sections:
        document.add_heading(f"{section['plan_name']} ({section['scope'].title()})", level=1)

        rows = [["Metric", "My Plan", "EvoFit AI Plan"]] + _metric_rows(section["comparison"])
        table = document.add_table(rows=len(rows), cols=3)
        table.style = "Light Grid Accent 1"
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                cell = table.cell(r, c)
                cell.text = str(value)
                if r == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True

        document.add_heading("AI Analysis", level=2)
        for obs in section["observations"]:
            document.add_paragraph(obs, style="List Bullet")

        if section["suggestions"]:
            document.add_heading("Suggestions", level=2)
            for s in section["suggestions"]:
                p = document.add_paragraph(style="List Bullet")
                run = p.add_run(f"{s['title']}: ")
                run.font.bold = True
                p.add_run(f"{s['detail']} ({s['reason']})")

        document.add_paragraph("")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
